from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/mac/.codex/attachments/13809a45-e0bc-47a8-a6f6-75e04d8b1f95/pasted-text.txt")
OUT_DIR = Path("/Users/mac/Desktop/欧美易经市场落地方案/outputs/hexagram_content")
MD_PATH = OUT_DIR / "dao_oracle_64_hexagrams_bilingual_database_ready.md"
JSONL_PATH = OUT_DIR / "dao_oracle_64_hexagrams_database_ready.jsonl"
DOCX_PATH = OUT_DIR / "dao_oracle_64_hexagrams_bilingual_微软雅黑.docx"
FONT_PRIMARY = "Microsoft YaHei"
FONT_CJK_FALLBACK = "PingFang SC"


SECTION_LABELS = [
    ("cosmic_timing", "COSMIC TIMING | 天时", "Cosmic Timing", "天时"),
    ("human_field", "HUMAN FIELD | 人和", "Human Field", "人和"),
    ("earthly_leverage", "EARTHLY LEVERAGE | 地利", "Earthly Leverage", "地利"),
    ("ancient_root", "ANCIENT ROOT | 卦脉溯源", "Ancient Root", "卦脉溯源"),
]

PROMPT_LINES = {
    "Why this energy is here",
    "Why this pause is happening",
    "Who surrounds this energy",
    "Who can help you stabilize",
    "What action is aligned",
    "What to do next",
    "Where this wisdom comes from",
    "何以有此能量",
    "何以有此停顿",
    "谁与此能量同行",
    "谁可助你稳定",
    "何以行之",
    "此慧何来",
}

STOP_PREFIXES = (
    "以上为",
    "可以按照",
    "继续",
    "好的，用户",
    "好的，我们继续",
    "好的，继续",
    "嗯，",
    "回顾之前",
    "现在需要",
    "我需要",
    "在撰写时",
    "好，思路",
    "接下来",
    "那么，",
    "用户",
    "开始逐一生成",
    "第一批",
    "第二批",
    "第三批",
    "第四批",
    "第五批",
    "第六批",
    "第七批",
    "第八批",
    "第九批",
    "第十批",
    "第十一批",
    "第十二批",
    "第十三批",
)

REPLACEMENTS = [
    ("This reading confirms that", "This reading suggests that"),
    ("This reading confirms", "This reading suggests"),
    ("此卦确认", "此卦提示"),
    ("It promises that", "It suggests that"),
    ("it promises that", "it suggests that"),
    ("It promises", "It suggests"),
    ("it promises", "it suggests"),
    ("它承诺", "它提示"),
    ("A moment of balance will come.", "A moment of balance can still emerge."),
    ("平衡的时刻会来的。", "平衡的时刻仍有可能出现。"),
    ("The rain will come.", "The rain is gathering."),
    ("雨会来的。", "雨正在聚集。"),
    ("You are not finished. And that is not a failure. It is the promise that the story continues, and you are still in it.",
     "You are not finished. And that is not a failure. It is a sign that the story continues, and you are still in it."),
    ("你尚未完成。而那不是失败。它是故事仍在继续、而你仍在其中的承诺。",
     "你尚未完成。而那不是失败。它是故事仍在继续、而你仍在其中的信号。"),
    ("The healing is not complete.", "The inner repair is not complete."),
    ("疗愈还没圆满。", "内在修复尚未完成。"),
    ("I am not yet healed from this.", "I am still recovering from this."),
    ("\"我还没从这个中痊愈。\"", "\"我仍在从这件事中恢复。\""),
    ("medicine for initial chaos", "remedy for initial chaos"),
    ("初始混沌唯一的解药", "初始混沌的重要解法"),
    ("diagnostic", "clarifying"),
    ("诊断性的", "澄清性的"),
    ("The bite breaks the skin so healing can begin.", "The decisive cut opens the way for repair to begin."),
    ("这一口咬破皮肉以便疗愈开始。", "这一次决断打开了修复开始的路径。"),
    ("gentleness is no longer the medicine", "gentleness is no longer the right tool"),
    ("温柔不再是解药", "温柔不再是合适的工具"),
    ("When the heart is pure and the intention is straight, the universe supports the movement.",
     "When the heart is clear and the intention is straight, the movement gains natural support."),
    ("当心纯净、意图直朴，宇宙支持其运动。",
     "当心清明、意图直朴，行动会获得自然的支撑。"),
    ("with the grain of the universe", "with the deeper grain of life"),
    ("与宇宙的纹理同向", "与生命更深的纹理同向"),
    ("the softest force in the universe", "the softest force in nature"),
    ("宇宙中最柔软的力量", "自然中最柔软的力量"),
    ("太阳依附于宇宙法则", "太阳依附于自然法则"),
    ("healing can begin", "repair can begin"),
    ("疗愈开始", "修复开始"),
    ("正在临在", "正在出现"),
    ("临在。", "出现。"),
    ("临在", "出现"),
]


def normalize(text: str) -> str:
    text = text.strip()
    text = text.replace("—", "—")
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text))


def dedupe_consecutive(lines: list[str]) -> list[str]:
    result: list[str] = []
    for line in lines:
        if result and result[-1] == line:
            continue
        result.append(line)
    return result


def paragraph_pairs(lines: list[str]) -> list[dict[str, str]]:
    cleaned: list[str] = []
    for line in lines:
        normalized = normalize(line)
        if not normalized or normalized in PROMPT_LINES:
            continue
        if normalized.startswith(STOP_PREFIXES):
            break
        cleaned.append(normalized)
    cleaned = dedupe_consecutive(cleaned)
    pairs: list[dict[str, str]] = []
    i = 0
    while i < len(cleaned):
        current = cleaned[i]
        nxt = cleaned[i + 1] if i + 1 < len(cleaned) else ""
        if current == "Traditional text:":
            pairs.append({"type": "label", "en": current, "zh": "传统经文："})
            i += 1
        elif not has_cjk(current) and nxt and has_cjk(nxt):
            pairs.append({"type": "pair", "en": current, "zh": nxt})
            i += 2
        elif has_cjk(current):
            pairs.append({"type": "zh", "en": "", "zh": current})
            i += 1
        else:
            pairs.append({"type": "en", "en": current, "zh": ""})
            i += 1
    return pairs


def find_top_level_starts(lines: list[str]) -> list[tuple[int, int, str]]:
    starts: list[tuple[int, int, str]] = []
    for i, line in enumerate(lines):
        match = re.match(r"^Hexagram (\d+):\s+(.+)$", line.strip())
        if not match:
            continue
        nonempty: list[str] = []
        for j in range(i + 1, min(i + 8, len(lines))):
            if lines[j].strip():
                nonempty.append(lines[j].strip())
        if nonempty and nonempty[0].startswith("卦") and (len(nonempty) < 2 or not nonempty[1].startswith("Upper trigram")):
            starts.append((i, int(match.group(1)), normalize(match.group(2))))
    return starts


def extract_records() -> list[dict]:
    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    starts = find_top_level_starts(raw_lines)

    blocks_by_number: dict[int, tuple[int, list[str], str]] = {}
    for idx, (start, number, title_en) in enumerate(starts):
        end = starts[idx + 1][0] if idx + 1 < len(starts) else len(raw_lines)
        # Keep the later duplicate for Hexagram 39; it is the corrected batch version.
        blocks_by_number[number] = (start, raw_lines[start:end], title_en)

    missing = [number for number in range(1, 65) if number not in blocks_by_number]
    if missing:
        raise RuntimeError(f"Missing hexagrams: {missing}")

    records: list[dict] = []
    for number in range(1, 65):
        _start, block, title_en = blocks_by_number[number]
        nonempty = [normalize(line) for line in block if normalize(line)]
        title_line = nonempty[0]
        zh_title = nonempty[1]
        core_en = nonempty[2]
        core_zh = nonempty[3]

        sections: dict[str, dict] = {}
        label_positions: list[tuple[int, str, str, str]] = []
        for key, label, display_en, display_zh in SECTION_LABELS:
            for idx, line in enumerate(block):
                if normalize(line) == label:
                    label_positions.append((idx, key, display_en, display_zh))
                    break
        if len(label_positions) != 4:
            raise RuntimeError(f"Hexagram {number} has {len(label_positions)} sections")
        label_positions.sort()

        for pos_idx, (pos, key, display_en, display_zh) in enumerate(label_positions):
            end = label_positions[pos_idx + 1][0] if pos_idx + 1 < len(label_positions) else len(block)
            body = block[pos + 1:end]
            sections[key] = {
                "title_en": display_en,
                "title_zh": display_zh,
                "paragraphs": paragraph_pairs(body),
            }

        records.append({
            "number": number,
            "title_en": normalize(title_line.replace(f"Hexagram {number}:", "").strip()),
            "title_zh": zh_title,
            "core_en": core_en,
            "core_zh": core_zh,
            "sections": sections,
            "question_lens": {
                "love": None,
                "career": None,
                "money": None,
                "wellbeing": None,
                "life_path": None,
            },
        })
    return records


def build_markdown(records: list[dict]) -> str:
    lines: list[str] = [
        "# Dao Oracle 64 Hexagrams - Bilingual Database-Ready Content",
        "",
        "整理原则：采用 64 卦完整结果卡结构，去除源文档中的内部推理和重复内容；统一为 70% 清晰实用 + 30% 神秘感；避免医疗诊断、绝对承诺和过度替宇宙下结论的表达。",
        "",
        "数据库字段建议：number, title_en, title_zh, core_en, core_zh, cosmic_timing, human_field, earthly_leverage, ancient_root, question_lens。",
        "",
    ]
    for record in records:
        lines.extend([
            f"## Hexagram {record['number']}: {record['title_en']}",
            "",
            record["title_zh"],
            "",
            f"**Core Message EN:** {record['core_en']}",
            "",
            f"**核心信息 ZH：** {record['core_zh']}",
            "",
        ])
        for key in ["cosmic_timing", "human_field", "earthly_leverage", "ancient_root"]:
            section = record["sections"][key]
            lines.extend([f"### {section['title_en']} | {section['title_zh']}", ""])
            for pair in section["paragraphs"]:
                if pair["en"]:
                    lines.append(pair["en"])
                if pair["zh"]:
                    lines.append(pair["zh"])
                lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT_PRIMARY
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_PRIMARY)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_PRIMARY)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_FALLBACK)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_pair_paragraph(doc: Document, en: str, zh: str, style: str | None = None) -> None:
    if en:
        p_en = doc.add_paragraph(style=style)
        run = p_en.add_run(en)
        set_run_font(run, size=9.5, color="1F2933")
        p_en.paragraph_format.space_after = Pt(2)
    if zh:
        p_zh = doc.add_paragraph(style=style)
        run = p_zh.add_run(zh)
        set_run_font(run, size=9.5, color="4B5563")
        p_zh.paragraph_format.space_after = Pt(7)


def build_docx(records: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = FONT_PRIMARY
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_PRIMARY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_PRIMARY)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_FALLBACK)
    styles["Normal"].font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Dao Oracle 64 Hexagrams")
    set_run_font(run, size=22, color="111827", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("中英文对照 · 网站调性优化版 · 数据库可录入结构")
    set_run_font(run, size=11, color="4B5563")

    doc.add_paragraph()
    summary = doc.add_paragraph()
    run = summary.add_run("整理说明")
    set_run_font(run, size=12, color="111827", bold=True)
    notes = [
        "已确认源文档覆盖 1-64 全部卦，第 39 卦有重复版本，本稿采用后面批量版本中的修正版。",
        "已去除源文档中的智能体推理、批次说明、重复确认语，只保留可发布结果卡内容。",
        "统一语言方向：70% 清晰实用 + 30% 神秘感；降低绝对承诺、医疗化表达和过度玄学判断。",
        "字体主声明为微软雅黑；因本机渲染环境无微软雅黑，中文设置了 PingFang SC 兜底，避免预览时丢字。",
        "JSONL 文件按 number/title/core/sections/question_lens 预留字段，便于后续数据库录入。",
    ]
    for note in notes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(note)
        set_run_font(run, size=9.5, color="374151")

    doc.add_section(WD_SECTION.NEW_PAGE)

    for record in records:
        h = doc.add_heading(level=1)
        run = h.add_run(f"Hexagram {record['number']}: {record['title_en']}")
        set_run_font(run, size=15, color="111827", bold=True)

        zh = doc.add_paragraph()
        run = zh.add_run(record["title_zh"])
        set_run_font(run, size=11, color="374151", bold=True)

        core_table = doc.add_table(rows=2, cols=1)
        core_table.style = "Table Grid"
        for cell in core_table._cells:
            set_cell_shading(cell, "F3F4F6")
        core_table.cell(0, 0).text = ""
        run = core_table.cell(0, 0).paragraphs[0].add_run(record["core_en"])
        set_run_font(run, size=10.5, color="111827", bold=True)
        core_table.cell(1, 0).text = ""
        run = core_table.cell(1, 0).paragraphs[0].add_run(record["core_zh"])
        set_run_font(run, size=10.5, color="374151", bold=True)
        doc.add_paragraph()

        for key in ["cosmic_timing", "human_field", "earthly_leverage", "ancient_root"]:
            section_data = record["sections"][key]
            h3 = doc.add_heading(level=2)
            run = h3.add_run(f"{section_data['title_en']} | {section_data['title_zh']}")
            set_run_font(run, size=12, color="111827", bold=True)
            for pair in section_data["paragraphs"]:
                if pair.get("type") == "label":
                    p = doc.add_paragraph()
                    run = p.add_run(f"{pair['en']} / {pair['zh']}")
                    set_run_font(run, size=9.5, color="111827", bold=True)
                else:
                    add_pair_paragraph(doc, pair["en"], pair["zh"])

        if record["number"] != 64:
            doc.add_page_break()

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    records = extract_records()

    MD_PATH.write_text(build_markdown(records), encoding="utf-8")
    with JSONL_PATH.open("w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    build_docx(records)

    print(f"records={len(records)}")
    print(MD_PATH)
    print(JSONL_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
