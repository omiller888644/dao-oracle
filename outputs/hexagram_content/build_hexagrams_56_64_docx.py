from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("/Users/mac/Desktop/欧美易经市场落地方案/outputs/hexagram_content")
MD_PATH = OUT_DIR / "hexagrams_56-64_bilingual_database_ready.md"
DOCX_PATH = OUT_DIR / "hexagrams_56-64_bilingual_database_ready_微软雅黑.docx"


HEXAGRAMS = [
    {
        "number": 56,
        "slug": "the-wanderer",
        "title_en": "The Wanderer",
        "title_zh": "旅",
        "pinyin": "Lu",
        "core_en": "You may be away from home, but you are not away from yourself.",
        "core_zh": "你也许离开了熟悉之地，但你并没有离开自己。",
        "sections": [
            ("Cosmic Timing", "天时",
             "A period of unfamiliarity is present. Hexagram 56 appears when you are in a place, role, relationship, or life stage where you cannot rely on old belonging. This is not punishment. It is an education in what remains steady when the surroundings change.",
             "一个陌生感强烈的阶段正在出现。旅卦常出现在你进入一个无法依赖旧归属的位置、关系、身份或人生阶段时。这不是惩罚，而是在训练你：当环境变化时，什么仍然属于你。"),
            ("Human Field", "人和",
             "As a wanderer, you may meet both helpful guides and cautious locals. Do not demand instant intimacy from people or spaces that do not yet know you. Look for the person who can orient you, even if they cannot fully welcome you.",
             "作为旅人，你会遇见愿意指路的人，也会遇见保持距离的人。不要要求陌生的人或环境立刻给你亲密感。先寻找那个能帮你辨认方向的人，即使他暂时不能完全接纳你。"),
            ("Earthly Leverage", "地利",
             "Your leverage is portable grounding. Choose one small practice, object, phrase, or rhythm that helps you feel centered anywhere. Do not compare the new place to the old one all day. Notice one useful or beautiful thing in the unfamiliar territory.",
             "你的现实杠杆，是可携带的稳定感。选择一个让你在任何地方都能回到中心的小练习、物品、短句或节律。不要整天拿新地方与旧地方比较。在陌生之地，先看见一个有用或美好的东西。"),
        ],
        "lens": {
            "Love": ("Do not ask a new relationship to feel like home too quickly. Let trust grow through small, consistent signs.",
                     "不要要求一段新关系太快像家一样安全。让信任通过小而稳定的信号慢慢长出来。"),
            "Career": ("You may be new to the room. Learn the terrain before trying to claim authority.",
                       "你可能刚进入一个新场域。先了解地形，再尝试建立权威。"),
            "Money": ("Avoid expensive decisions made to soothe displacement. Stabilize basics before expanding.",
                      "不要为了缓解漂泊感而做高成本决定。先稳定基本盘，再谈扩张。"),
            "Wellbeing": ("Loneliness may be present. Use routine, movement, and simple comforts to help the body feel located.",
                          "孤独感可能存在。用规律、身体活动和简单安慰，让身体重新感到有位置。"),
            "Life Path": ("This season may not give you permanence, but it can teach you what travels with you.",
                          "这个阶段未必给你稳定归宿，但会教你看清：什么会一直跟随你。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☲ Li, Flame. Lower trigram: ☶ Gen, Mountain.",
            "trigrams_zh": "上卦：☲ 离，火。下卦：☶ 艮，山。",
            "text": "旅，小亨，旅贞吉。",
            "note_en": "The Wanderer teaches small success, humility, and correct conduct in unfamiliar territory.",
            "note_zh": "旅卦教导人在陌生处境中取得小亨通：谦逊、守正、懂得分寸。",
        },
    },
    {
        "number": 57,
        "slug": "the-gentle",
        "title_en": "The Gentle",
        "title_zh": "巽",
        "pinyin": "Xun",
        "core_en": "Gentleness is not weakness. It is influence that knows how to enter without breaking.",
        "core_zh": "温柔不是软弱，而是一种懂得进入、却不破坏的影响力。",
        "sections": [
            ("Cosmic Timing", "天时",
             "Subtle influence is stronger than force right now. Hexagram 57 appears when direct pressure would create resistance, but steady presence can reach the place that argument cannot.",
             "此刻，微妙影响比强力推进更有效。巽卦常出现在直接施压会引发抗拒、但稳定在场能抵达争论无法抵达之处的时候。"),
            ("Human Field", "人和",
             "Someone may be resistant to a direct approach. Do not confront the closed door with more force. Use a question, a demonstration, a quieter tone, or repeated consistency. Also notice the subtle ally whose influence is quiet but real.",
             "某个人可能抗拒直接方式。不要用更大的力量撞门。试着用一个问题、一次示范、更安静的语气，或持续一致的行动。同时留意那个影响力安静却真实的盟友。"),
            ("Earthly Leverage", "地利",
             "Your leverage is gentle persistence. Choose one situation where pushing has failed. This week, try the indirect path: less explanation, more example; less demand, more rhythm.",
             "你的现实杠杆，是温柔而持续的坚持。选一个你越推越没结果的情况。本周试试间接路径：少解释，多示范；少要求，多节律。"),
        ],
        "lens": {
            "Love": ("Do not force emotional disclosure. Create safety through consistency and a softer invitation.",
                     "不要强迫对方袒露情感。用稳定和柔和邀请创造安全感。"),
            "Career": ("Influence through reliability, not volume. Let your work become the argument.",
                       "用可靠性影响别人，而不是用音量。让你的工作本身成为论点。"),
            "Money": ("Small repeated improvements matter more than one dramatic move.",
                      "微小但重复的改善，比一次戏剧性动作更重要。"),
            "Wellbeing": ("Use gentle repetition: breath, stretching, hydration, and rhythm. Do not punish the body into change.",
                          "使用温和重复：呼吸、拉伸、补水、节律。不要用惩罚身体的方式逼迫改变。"),
            "Life Path": ("The next door may open through subtle alignment, not through conquest.",
                          "下一扇门可能不是靠征服打开，而是靠微妙对齐打开。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☴ Xun, Wind. Lower trigram: ☴ Xun, Wind.",
            "trigrams_zh": "上卦：☴ 巽，风。下卦：☴ 巽，风。",
            "text": "巽，小亨。利有攸往，利见大人。",
            "note_en": "The Gentle teaches penetration through consistency, adaptation, and quiet influence.",
            "note_zh": "巽卦教导以恒常、适应和安静影响来渗透，而非以强攻取胜。",
        },
    },
    {
        "number": 58,
        "slug": "the-joyous",
        "title_en": "The Joyous",
        "title_zh": "兑",
        "pinyin": "Dui",
        "core_en": "Joy is not proof that life is perfect. It is the water returning to its natural brightness.",
        "core_zh": "喜悦不是生活完美的证明，而是水回到本来明亮的状态。",
        "sections": [
            ("Cosmic Timing", "天时",
             "A moment of shared joy is available. Hexagram 58 does not claim that every problem is solved. It says joy can still be reached, especially through openness, conversation, and sincere delight.",
             "一个可以共享喜悦的时刻正在出现。兑卦并不是说所有问题都解决了，而是说：喜悦仍然可以被触及，尤其通过敞开、交流和真诚的愉悦。"),
            ("Human Field", "人和",
             "Joy expands when it is shared with the right people. Move toward those who can laugh with you without making you perform. Be careful with people who drain every bright moment through cynicism or silence.",
             "喜悦在与对的人分享时会扩大。靠近那些能与你自然欢笑、不会要求你表演的人。也要留意那些用冷嘲或沉默耗掉每个明亮时刻的人。"),
            ("Earthly Leverage", "地利",
             "Your leverage is the deliberate cultivation of delight. Do one small thing purely because it brings life back into the room. Share one simple joy as an offering, not as a performance.",
             "你的现实杠杆，是有意识地培育喜悦。做一件纯粹让生命感回到房间里的小事。分享一个简单的快乐，把它当作赠予，而不是表演。"),
        ],
        "lens": {
            "Love": ("Look for ease, laughter, and emotional generosity. Joy should not require self-betrayal.",
                     "看见轻松、笑声和情感慷慨。真正的喜悦不需要你背叛自己。"),
            "Career": ("Team morale matters. A lighter exchange may unlock more progress than another serious meeting.",
                       "团队士气很重要。一次轻松交流，可能比又一场严肃会议更能推动事情。"),
            "Money": ("Enjoyment is allowed, but keep it conscious. Spend on what genuinely nourishes, not what briefly distracts.",
                      "享受是可以的，但要有意识。把钱花在真正滋养你的东西上，而不是短暂麻痹你的东西上。"),
            "Wellbeing": ("Notice small pleasures that regulate the body: warmth, music, taste, conversation, sunlight.",
                          "留意能调节身体的小愉悦：温暖、音乐、味道、对话、阳光。"),
            "Life Path": ("Joy is information. What genuinely delights you may be pointing toward aliveness.",
                          "喜悦也是信息。真正让你明亮的东西，可能正在指向生命力。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☱ Dui, Lake. Lower trigram: ☱ Dui, Lake.",
            "trigrams_zh": "上卦：☱ 兑，泽。下卦：☱ 兑，泽。",
            "text": "兑，亨，利贞。",
            "note_en": "The Joyous teaches that delight becomes auspicious when it remains sincere and rightly held.",
            "note_zh": "兑卦教导：喜悦若保持真诚并守正，就能通达而有益。",
        },
    },
    {
        "number": 59,
        "slug": "dispersion",
        "title_en": "Dispersion",
        "title_zh": "涣",
        "pinyin": "Huan",
        "core_en": "What has frozen must flow again. Release is not loss; it is movement returning.",
        "core_zh": "冻结之物需要重新流动。释放不是失去，而是运动的回归。",
        "sections": [
            ("Cosmic Timing", "天时",
             "Something rigid is loosening: a belief, resentment, role, identity, or structure. Hexagram 59 appears when the old container can no longer hold life well. The ice is melting so circulation can return.",
             "某个僵硬之物正在松动：一个信念、一段怨气、一个角色、身份或结构。涣卦常出现在旧容器已无法好好承载生命的时候。冰正在融化，是为了让循环回来。"),
            ("Human Field", "人和",
             "You may be separating from people or roles that once defined you. This can feel like loss, but it can also be liberation. Move toward those who do not need you to stay frozen for their comfort.",
             "你可能正在离开曾经定义你的人或角色。这会像失去，但也可能是解放。靠近那些不需要你为了他们的舒适而继续冻结的人。"),
            ("Earthly Leverage", "地利",
             "Your leverage is release with a container. Name one rigid thing you are holding. Ask what would happen if it softened. If you feel scattered, use one small bank: a schedule, a boundary, a ritual, or a clear next step.",
             "你的现实杠杆，是带着容器的释放。说出一个你紧紧抓住的僵硬之物，然后问：如果它变软，会发生什么？若你感到散乱，就使用一个小小的河岸：时间表、边界、仪式或清晰的下一步。"),
        ],
        "lens": {
            "Love": ("A fixed story about someone may need to dissolve. Let the relationship breathe before deciding its final meaning.",
                     "你对某个人的固定故事可能需要松开。在决定关系的最终意义之前，先让它重新呼吸。"),
            "Career": ("Old structures may be breaking down. Let unnecessary roles, meetings, or assumptions disperse.",
                       "旧结构可能正在瓦解。让不必要的角色、会议或假设散开。"),
            "Money": ("Release financial patterns that keep you stuck. Create flow, but keep boundaries.",
                      "释放让你卡住的财务模式。创造流动，但保留边界。"),
            "Wellbeing": ("The body may need movement, water, breath, or emotional release. Do not confuse release with collapse.",
                          "身体可能需要运动、水、呼吸或情绪释放。不要把释放误认为崩溃。"),
            "Life Path": ("An old identity may be dissolving. This is not the end of you; it is space returning.",
                          "一个旧身份可能正在溶解。这不是你的结束，而是空间正在回来。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☴ Xun, Wind. Lower trigram: ☵ Kan, Water.",
            "trigrams_zh": "上卦：☴ 巽，风。下卦：☵ 坎，水。",
            "text": "涣，亨。王假有庙。利涉大川，利贞。",
            "note_en": "Dispersion teaches sacred release: what loosens may return to circulation rather than disappear.",
            "note_zh": "涣卦教导神圣的释放：松开的东西未必消失，它可能只是重新进入循环。",
        },
    },
    {
        "number": 60,
        "slug": "limitation",
        "title_en": "Limitation",
        "title_zh": "节",
        "pinyin": "Jie",
        "core_en": "Limits are not walls. They are banks that give the river power.",
        "core_zh": "限制不是墙，而是让河流有力量的岸。",
        "sections": [
            ("Cosmic Timing", "天时",
             "A call for proportion is present. Hexagram 60 appears when something in life needs form, measure, or a clean boundary. This is not deprivation. It is containment that makes movement possible.",
             "一个关于分寸的召唤正在出现。节卦常出现在生活中某件事需要形式、尺��或清晰边界的时候。这不是剥夺，而是让流动成为可能的容器。"),
            ("Human Field", "人和",
             "You may need to set a limit with someone, or respect a limit someone else has set. Healthy limits are definitions, not punishments. They say: within this, life can be held.",
             "你可能需要对某人设限，或尊重他人设下的限制。健康边界是定义，不是惩罚。它说的是：在这个范围内，生命可以被承载。"),
            ("Earthly Leverage", "地利",
             "Your leverage is the boundary that protects the yes. Choose one area where energy is leaking: time, spending, attention, emotional labor. Install one clear limit this week.",
             "你的现实杠杆，是保护“是”的那个边界。选择一个能量正在泄漏的领域：时间、消费、注意力、情感劳动。本周设立一个清晰限制。"),
        ],
        "lens": {
            "Love": ("A loving boundary is still a boundary. Closeness becomes safer when expectations are clear.",
                     "有爱的边界仍然是边界。当期待清晰时，亲密才更安全。"),
            "Career": ("Define scope, deadline, ownership, and what is not included. Vague work drains power.",
                       "明确范围、截止期、负责人，以及不包含什么。模糊的工作会耗尽力量。"),
            "Money": ("Budgeting is not punishment. It is a riverbank for your resources.",
                      "预算不是惩罚，而是你资源的河岸。"),
            "Wellbeing": ("Limits around sleep, screens, food, or commitments may restore energy faster than another breakthrough attempt.",
                          "围绕睡眠、屏幕、饮食或承诺设限，可能比追求另一次突破更快恢复能量。"),
            "Life Path": ("Freedom without form becomes diffusion. Choose the limits that help your life deepen.",
                          "没有形式的自由会变成涣散。选择那些能让人生变深的限制。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☵ Kan, Water. Lower trigram: ☱ Dui, Lake.",
            "trigrams_zh": "上卦：☵ 坎，水。下卦：☱ 兑，泽。",
            "text": "节，亨。苦节不可贞。",
            "note_en": "Limitation teaches that boundaries are necessary, but harsh restriction cannot be sustained.",
            "note_zh": "节卦教导边界是必要的，但过度苦涩的限制不可长久持守。",
        },
    },
    {
        "number": 61,
        "slug": "inner-truth",
        "title_en": "Inner Truth",
        "title_zh": "中孚",
        "pinyin": "Zhong Fu",
        "core_en": "Truth does not need to shout. It resonates.",
        "core_zh": "真相不需要喊叫。它会产生共鸣。",
        "sections": [
            ("Cosmic Timing", "天时",
             "A moment of inner sincerity is present. Hexagram 61 appears when the performed self is less useful than the genuine self. What is true may not need more proof; it needs a cleaner channel.",
             "一个内在真诚的时刻正在出现。中孚卦常出现在“表演出来的自己”不如“真实的自己”有力量的时候。真正的东西未必需要更多证明，而是需要更清澈的通道。"),
            ("Human Field", "人和",
             "People can feel sincerity even before they can explain it. Those aligned with you may recognize your truth without pressure. Those not aligned may resist no matter how much evidence you bring.",
             "人们常在说清楚之前就感知到真诚。与你对齐的人，不需要被压迫也能认出你的真实；不对齐的人，即使你给再多证据，也可能抗拒。"),
            ("Earthly Leverage", "地利",
             "Your leverage is sincere expression. In one important conversation, speak from actual experience rather than from what you think you should say. Reduce performance. Increase truth.",
             "你的现实杠杆，是真诚表达。在一次重要对话中，从真实经验出发，而不是从“我应该怎么说”出发。减少表演，增加真实。"),
        ],
        "lens": {
            "Love": ("Say the true thing gently. A bond that can hold sincerity becomes stronger.",
                     "温柔地说出真实。能承载真诚的关系，会因此更强。"),
            "Career": ("Credibility grows when your words, work, and motives align.",
                       "当你的语言、工作和动机一致时，可信度会增长。"),
            "Money": ("Be honest about what you value and what you fear. Financial clarity begins with inner clarity.",
                      "诚实面对你重视什么、害怕什么。财务清晰始于内在清晰。"),
            "Wellbeing": ("Notice where you are performing wellness instead of actually feeling well.",
                          "留意你在哪里是在表演“我很好”，而不是真正感到安稳。"),
            "Life Path": ("Your direction becomes clearer when you stop negotiating against your own truth.",
                          "当你停止与自己的真实讨价还价时，方向会变清楚。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☴ Xun, Wind. Lower trigram: ☱ Dui, Lake.",
            "trigrams_zh": "上卦：☴ 巽，风。下卦：☱ 兑，泽。",
            "text": "中孚，豚鱼吉。利涉大川，利贞。",
            "note_en": "Inner Truth teaches influence through sincerity: even simple creatures respond to what is genuine.",
            "note_zh": "中孚卦教导以真诚产生影响：真实之物连朴素生命也能感知。",
        },
    },
    {
        "number": 62,
        "slug": "small-excess",
        "title_en": "Small Excess",
        "title_zh": "小过",
        "pinyin": "Xiao Guo",
        "core_en": "A small correction, made with care, can change the whole direction.",
        "core_zh": "一个被认真完成的小修正，可能改变整个方向。",
        "sections": [
            ("Cosmic Timing", "天时",
             "The situation does not need a grand gesture. It needs precision. Hexagram 62 appears when the big picture may be stable, but a small imbalance is asking for attention.",
             "此刻不需要宏大姿态，而需要精准。小过卦常出现在大局尚可稳定、但某个小失衡正在要求被看见的时候。"),
            ("Human Field", "人和",
             "Do not overlook the small steady people in your life. The person who remembers details, checks in quietly, or offers modest help may be more important than the dramatic ally you imagined.",
             "不要忽略你生活里那些小而稳定的人。那个记得细节、安静问候、提供小帮助的人，可能比你想象中的戏剧性盟友更重要。"),
            ("Earthly Leverage", "地利",
             "Your leverage is the smallest useful action. Finish a small neglected task, repair a minor misunderstanding, or adjust one habit. Do not overcorrect. Fine-tune.",
             "你的现实杠杆，是最小但有用的行动。完成一件被忽略的小事，修复一个小误会，或调整一个习惯。不要过度修正，微调即可。"),
        ],
        "lens": {
            "Love": ("A small apology, check-in, or changed tone may matter more than a dramatic confession.",
                     "一个小小的道歉、问候或语气改变，可能比戏剧性告白更重要。"),
            "Career": ("Fix the small workflow issue before attempting a major overhaul.",
                       "先修复小流程问题，再考虑大规模改革。"),
            "Money": ("Small expenses, missed payments, or tiny leaks deserve attention now.",
                      "小支出、漏缴款项或微小漏洞，此刻值得被认真处理。"),
            "Wellbeing": ("One small body habit may be the hinge: hydration, posture, sleep time, breathing.",
                          "一个小身体习惯可能是关键：补水、姿势、睡眠时间、呼吸。"),
            "Life Path": ("Do not wait for a grand sign. The next direction may arrive as a small adjustment.",
                          "不要等待宏大征兆。下一步方向可能以一个小调整的形式出现。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☳ Zhen, Thunder. Lower trigram: ☶ Gen, Mountain.",
            "trigrams_zh": "上卦：☳ 震，雷。下卦：☶ 艮，山。",
            "text": "小过，亨，利贞。可小事，不可大事。",
            "note_en": "Small Excess teaches that modest, careful action is favored over grand ambition.",
            "note_zh": "小过卦教导：此时小事可为，大事不宜，谨慎微调胜过宏大冒进。",
        },
    },
    {
        "number": 63,
        "slug": "after-completion",
        "title_en": "After Completion",
        "title_zh": "既济",
        "pinyin": "Ji Ji",
        "core_en": "Completion is real, but it is not permanent unless it is tended.",
        "core_zh": "完成是真实的，但若无人照看，它不会自动长久。",
        "sections": [
            ("Cosmic Timing", "天时",
             "A cycle has reached completion. Hexagram 63 marks a rare moment of balance, but it also carries a warning: what has been completed can unravel if neglected.",
             "一个循环已经抵达完成。既济卦标记一种罕见的平衡，但也带着提醒：完成之物若被忽略，仍可能松散。"),
            ("Human Field", "人和",
             "Some people may drift away after a shared goal is achieved. Others will remain for the next chapter. Let the sorting happen. Celebrate with those who can meet your success without envy.",
             "共同目标达成后，有些人可能会离开，有些人会留下进入下一章。让自然筛选发生。与那些能不带嫉妒地见证你成功的人庆祝。"),
            ("Earthly Leverage", "地利",
             "Your leverage is celebration plus maintenance. Acknowledge what is finished. Then identify what must be tended so the achieved order does not decay.",
             "你的现实杠杆，是庆祝加维护。承认你已经完成的事，然后找出需要继续照看的部分，让已经形成的秩序不至于衰败。"),
        ],
        "lens": {
            "Love": ("A relationship milestone is not the end of care. Keep tending the bond after the moment of success.",
                     "关系里的里程碑不是照顾的终点。成功时刻之后，仍要继续照看这段联结。"),
            "Career": ("Delivery is not the end. Document, stabilize, hand off, and protect the standard.",
                       "交付不是结束。记录、稳定、交接，并保护已建立的标准。"),
            "Money": ("A financial goal may be reached. Now create the habit that preserves it.",
                      "一个财务目标可能已经达成。现在要建立能守住它的习惯。"),
            "Wellbeing": ("Progress needs maintenance. Do not abandon the practices that helped you arrive here.",
                          "进展需要维护。不要放弃那些帮助你走到这里的练习。"),
            "Life Path": ("Honor the completed chapter, then prepare for the next cycle with humility.",
                          "尊重已经完成的章节，然后带着谦逊准备下一个循环。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☵ Kan, Water. Lower trigram: ☲ Li, Flame.",
            "trigrams_zh": "上卦：☵ 坎，水。下卦：☲ 离，火。",
            "text": "既济，亨小，利贞。初吉终乱。",
            "note_en": "After Completion teaches that balance is precious but fragile. Completion must be maintained.",
            "note_zh": "既济卦教导：平衡珍贵但脆弱。完成之后，仍需维护。",
        },
    },
    {
        "number": 64,
        "slug": "before-completion",
        "title_en": "Before Completion",
        "title_zh": "未济",
        "pinyin": "Wei Ji",
        "core_en": "You are not finished. That is not failure; it is the proof that life is still moving.",
        "core_zh": "你尚未完成。这不是失败，而是生命仍在运动的证明。",
        "sections": [
            ("Cosmic Timing", "天时",
             "Something is not yet complete. The final piece has not landed, the answer has not fully arrived, or the crossing is not finished. Hexagram 64 closes the I Ching by refusing closure.",
             "某件事尚未完成。最后一块还没有落位，答案还没有完全到来，渡河还没有结束。未济作为易经最后一卦，恰恰以拒绝封口来结束。"),
            ("Human Field", "人和",
             "Do not compare your unfinished middle with someone else's polished ending. The right companions are not the people who demand completion from you, but those who can walk with your becoming.",
             "不要拿自己未完成的中段，去比较别人打磨好的结尾。真正适合同行的人，不是催你完成的人，而是能陪你继续成为的人。"),
            ("Earthly Leverage", "地利",
             "Your leverage is staying in the process. Name what is unfinished without shame. Take one step, not to force the ending, but to continue the crossing with care.",
             "你的现实杠杆，是愿意留在过程中。不带羞耻地说出什么尚未完成。迈出一步，不是为了逼出结局，而是为了更谨慎地继续渡河。"),
        ],
        "lens": {
            "Love": ("Do not demand final certainty from a bond still forming. Watch whether both people keep crossing.",
                     "不要向仍在形成的关系索要最终确定。观察双方是否都还在继续渡河。"),
            "Career": ("The project is not mature yet. Protect the last stage from rushing, distraction, or premature celebration.",
                       "项目尚未成熟。保护最后阶段，避免仓促、分心或过早庆祝。"),
            "Money": ("Do not count the outcome before it lands. Keep reserves and finish the last details.",
                      "结果尚未落地前，不要提前计算胜利。保留余地，完成最后细节。"),
            "Wellbeing": ("You may be in the middle of recovery or adjustment. Respect the process; avoid declaring failure too early.",
                          "你可能正处在恢复或调整中段。尊重过程，不要太早宣布失败。"),
            "Life Path": ("The story continues. Your unfinished state may be the most honest place from which to choose the next step.",
                          "故事仍在继续。你的未完成状态，可能正是选择下一步最诚实的位置。"),
        },
        "root": {
            "trigrams_en": "Upper trigram: ☲ Li, Flame. Lower trigram: ☵ Kan, Water.",
            "trigrams_zh": "上卦：☲ 离，火。下卦：☵ 坎，水。",
            "text": "未济，亨。小狐汔济，濡其尾，无攸利。",
            "note_en": "Before Completion teaches that the almost-finished moment requires care. The ending is near, but not yet secure.",
            "note_zh": "未济卦教导：接近完成之时最需要谨慎。终点已近，��尚未稳固。",
        },
    },
]


def md_escape(text: str) -> str:
    return text.replace("\n", " ")


def build_markdown() -> str:
    lines = [
        "# Dao Oracle Revised Bilingual Content: Hexagrams 56-64",
        "",
        "> 用途：第 56-64 卦内容库样稿，可用于网站前台展示、Word 审稿和后期数据库录入。",
        "> 格式：英文前台文案 + 中文理解逐段对照。",
        "> 调性：70% clear and practical, 30% mystical and ritual.",
        "> 数据库建议：按 hexagram_number、section_key、language、question_lens 分字段入库。",
        "",
        "---",
        "",
    ]
    for h in HEXAGRAMS:
        lines += [
            f"# Hexagram {h['number']}: {h['title_zh']} {h['pinyin']}",
            "",
            f"## {h['title_en']} / {h['title_zh']}",
            "",
            f"**Core Message**  ",
            h["core_en"],
            "",
            f"**核心提示**  ",
            h["core_zh"],
            "",
        ]
        for key_en, key_zh, en, zh in h["sections"]:
            lines += [
                f"### {key_en} / {key_zh}",
                "",
                en,
                "",
                zh,
                "",
            ]
        lines += ["### Question Lens / 问题方向解读", ""]
        label_zh = {
            "Love": "爱情",
            "Career": "事业",
            "Money": "金钱",
            "Wellbeing": "身心状态",
            "Life Path": "人生方向",
        }
        for label, (en, zh) in h["lens"].items():
            lines += [f"**{label} / {label_zh[label]}**  ", en, "", zh, ""]
        root = h["root"]
        lines += [
            "### Ancient Root / 卦脉溯源",
            "",
            f"Hexagram {h['number']}: {h['title_zh']} {h['pinyin']} — {h['title_en']}",
            "",
            root["trigrams_en"],
            "",
            root["trigrams_zh"],
            "",
            "Traditional text / 传统经文：",
            "",
            f"> {root['text']}",
            "",
            root["note_en"],
            "",
            root["note_zh"],
            "",
            "---",
            "",
        ]
    return "\n".join(lines)


def set_run_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.2):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(doc, text, style=None, bold=False, italic=False, color=None, size=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return p


def setup_styles(doc):
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(34, 34, 34)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    set_paragraph_spacing(p, before=12 if level == 1 else 8, after=6)
    run = p.add_run(text)
    set_run_font(run)
    run.bold = True
    return p


def add_lens_table(doc, lens):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Direction", "English", "中文理解"]):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r)
        r.bold = True
    label_zh = {
        "Love": "爱情",
        "Career": "事业",
        "Money": "金钱",
        "Wellbeing": "身心状态",
        "Life Path": "人生方向",
    }
    for label, (en, zh) in lens.items():
        cells = table.add_row().cells
        values = [f"{label}\n{label_zh[label]}", en, zh]
        for cell, value in zip(cells, values):
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=2, line=1.15)
            r = p.add_run(value)
            set_run_font(r)
            r.font.size = Pt(9.5)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=3)
    r = title.add_run("Dao Oracle 卦文内容库：第56-64卦")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=14)
    r = subtitle.add_run("中英文对照 · 网站调性修订版 · 数据库录入友好")
    set_run_font(r)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "调整说明", 1)
    notes = [
        "已删除原稿中另一个智能体的推理/自述，只保留可发布内容。",
        "语言从长篇灵性散文调整为结果卡片：清晰实用为主，保留东方宇宙仪式感。",
        "每卦增加 Question Lens，便于根据用户问题方向调用不同解释。",
        "健康相关统一改为 Wellbeing / 身心状态，避免医疗诊断和承诺。",
        "Word 字体统一使用微软雅黑；Markdown 版本可作为数据库录入参考。",
    ]
    for note in notes:
        p = doc.add_paragraph(style=None)
        set_paragraph_spacing(p, after=3)
        r = p.add_run("• " + note)
        set_run_font(r)
        r.font.size = Pt(10)

    for h in HEXAGRAMS:
        doc.add_page_break()
        add_heading(doc, f"Hexagram {h['number']}: {h['title_zh']} {h['pinyin']} / {h['title_en']}", 1)
        add_para(doc, "Core Message / 核心提示", bold=True, color="1F4D78", before=2, after=3)
        add_para(doc, h["core_en"], italic=True, size=11, after=3)
        add_para(doc, h["core_zh"], size=10.5, after=8)
        for key_en, key_zh, en, zh in h["sections"]:
            add_heading(doc, f"{key_en} / {key_zh}", 2)
            add_para(doc, en, size=10.5, after=3)
            add_para(doc, zh, size=10.5, after=8)
        add_heading(doc, "Question Lens / 问题方向解读", 2)
        add_lens_table(doc, h["lens"])
        add_heading(doc, "Ancient Root / 卦脉溯源", 2)
        root = h["root"]
        add_para(doc, f"Hexagram {h['number']}: {h['title_zh']} {h['pinyin']} — {h['title_en']}", bold=True, after=3)
        add_para(doc, root["trigrams_en"], after=2)
        add_para(doc, root["trigrams_zh"], after=6)
        add_para(doc, f"Traditional text / 传统经文：{root['text']}", italic=True, after=6)
        add_para(doc, root["note_en"], after=3)
        add_para(doc, root["note_zh"], after=6)

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
