from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/mac/Desktop/欧美易经市场落地方案/outputs/brand_doc/品牌命名与解卦结构方案.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        table._tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    table._tbl.insert(1, tbl_grid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="F2F4F7"):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(31, 41, 55)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(17, 24, 39)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
    else:
        r2 = p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_callout(doc, label, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(text)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(31, 41, 55)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("品牌命名与解卦结构方案")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(22)
r.bold = True
r.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
r = subtitle.add_run("Dao Oracle / 欧美易经产品定位、Slogan 与结果页框架")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(85, 85, 85)

add_callout(
    doc,
    "核心结论",
    "推荐品牌名为 Dao Oracle，Slogan 为 Ancient Eastern wisdom for modern life decisions. "
    "解卦结果不应只输出卦文和传统解释，而应采用“天时 / 人和 / 地利”的三层结构，把东方易经逻辑翻译成欧美用户能理解的个人洞察、关系影响和行动条件。",
    fill="E8EEF5",
)

add_heading(doc, "1. 市场判断：欧美用户买的不是“算命”，而是清晰感", 1)
add_body(
    doc,
    "欧美主流身心灵产品，如 Co-Star、The Pattern、Labyrinthos、Trusted Tarot 和冥想类 wellness app，解决的核心问题不是传统意义上的预测未来，而是帮助用户理解自己、解释关系、缓解焦虑、做出下一步选择。"
)
add_body(
    doc,
    "这与我们的目标用户高度一致：用户在关系、职业、人生选择和情绪不安中，需要一个有仪式感的暂停，一个象征系统，以及一个可执行的行动建议。"
)

table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
for i, text in enumerate(["产品类型", "用户为什么用", "解决的问题", "对我们的启发"]):
    hdr[i].text = text
rows = [
    ["Co-Star / The Pattern", "每日状态、关系洞察、人格模式", "被理解、关系解释、阶段判断", "结果要像“懂我”，不能是泛泛祝福"],
    ["Labyrinthos / Tarot", "学习古老系统，获得每日智慧", "仪式感、象征系统、可收藏内容", "易经可包装成 64 种能量原型"],
    ["Trusted Tarot", "快速获得免费/低价读牌", "低门槛、高频次、每日指引", "首次起卦免费，深度解读再付费"],
    ["Mindfulness / wellness apps", "减压、稳定情绪、建立意义感", "情绪安定、日常练习、精神支持", "结果页要给 ritual / action，而不只是解释"],
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
set_table_width(table, [1900, 2300, 2400, 2760])
style_table(table)

add_heading(doc, "2. 品牌命名建议", 1)
add_body(doc, "第一推荐：Dao Oracle", bold_prefix="第一推荐：")
add_body(doc, "Slogan: Ancient Eastern wisdom for modern life decisions.")
add_body(doc, "中文理解：用古老东方智慧，回应现代人生抉择。")
add_body(
    doc,
    "Dao 保留东方哲学根基，但比 Yijing / Zhouyi 更容易发音和理解。Oracle 是欧美用户熟悉的“神谕 / 指引”语义，能让用户马上知道这是一个提供洞察和方向的工具。"
)
add_bullet(doc, "主品牌：Dao Oracle")
add_bullet(doc, "品牌补充语：Ancient Eastern wisdom for modern life decisions.")
add_bullet(doc, "首页按钮：Begin Your Reading")
add_bullet(doc, "一句话解释：Ask your question. Cast the ancient 64 signs. Receive personal guidance for clarity, change, and alignment.")

add_heading(doc, "3. 为什么不能只输出卦文和解卦", 1)
add_body(
    doc,
    "如果结果页只显示卦文、卦象名称和传统解释，欧美用户会遇到三个问题：一是不知道卦文和自己问题的关系；二是不知道周围人和现实条件如何参与这件事；三是不知道下一步该做什么。"
)
add_body(
    doc,
    "因此，传统卦文应该作为底层依据和高级内容，而不是首屏的主体。首屏应先翻译成“我现在处在什么能量状态、关系中发生了什么、我可以借助什么条件行动”。"
)

add_callout(
    doc,
    "产品原则",
    "前端不直接卖“卦文”，而是卖“把卦象转化为现代人生处境的解释能力”。这才是欧美用户愿意付费的地方。",
    fill="FFF4CC",
)

add_heading(doc, "4. “天时 / 人和 / 地利”解卦结构", 1)
add_body(
    doc,
    "这个结构非常适合欧美市场，因为它把东方哲学中的整体观转译为用户能理解的三类问题：我现在处在什么状态？周围人的角色是什么？我能利用哪些现实条件跳出困境？"
)

table = doc.add_table(rows=1, cols=4)
for i, text in enumerate(["中文模块", "英文前端命名", "回答的问题", "内容重点"]):
    table.rows[0].cells[i].text = text
structure_rows = [
    ["天时", "Cosmic Timing", "宇宙/变化现在给我的回应是什么？", "能量状态、当前阻滞、为什么这件事现在发生、卦象把用户带入的核心困境"],
    ["人和", "Human Field", "我和周围人的关系如何影响这件事？", "伴侣、朋友、同事、家人、客户等角色；谁推动、谁阻碍、谁是镜子、谁是变量"],
    ["地利", "Earthly Leverage", "我可以利用哪些现实条件走出来？", "资源、环境、时机、信息、渠道、身体状态、可执行动作和短期策略"],
]
for row in structure_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
set_table_width(table, [1100, 1800, 2800, 3660])
style_table(table, header_fill="E8EEF5")

add_heading(doc, "5. 建议的结果页输出模板", 1)
add_body(doc, "推荐结果页不直接以“Hexagram 1: Qian”开场，而是用品牌化能量名和三层解读开场。")

add_body(doc, "结果页结构：")
for item in [
    "Your Oracle: Universal Flow / Deep Grounding / The Turning Point 等品牌化能量名",
    "Cosmic Timing: The universe is pointing to the energy pattern behind your question.",
    "Human Field: The people around you are shaping the situation in these roles.",
    "Earthly Leverage: The resources and conditions you can use now.",
    "Ritual Action: One small action within 24 hours.",
    "Ancient Root: 可折叠显示传统卦名、卦文和 I Ching 依据。",
]:
    add_bullet(doc, item)

add_callout(
    doc,
    "付费边界建议",
    "免费版给出能量名、简短 Cosmic Timing 和一句 Ritual Action；付费版解锁完整三层分析、人际角色拆解、可执行策略、传统卦文依据和分享卡片。",
    fill="E8EEF5",
)

add_heading(doc, "6. AI Prompt 输出字段", 1)
table = doc.add_table(rows=1, cols=3)
for i, text in enumerate(["字段", "用途", "示例内容"]):
    table.rows[0].cells[i].text = text
prompt_rows = [
    ["energy_name", "品牌化卦象名", "Universal Flow"],
    ["cosmic_timing", "天时 / Cosmic Timing", "You are being asked to move, but not from panic. The timing favors directed action."],
    ["human_field", "人和 / Human Field", "One person may be amplifying your doubt; another is offering a quiet opening."],
    ["earthly_leverage", "地利 / Earthly Leverage", "Use the information already in your hands. Do not wait for perfect certainty."],
    ["ritual_action", "24 小时行动建议", "Write one message you have avoided, but do not send it for one hour."],
    ["ancient_root", "传统依据", "Hexagram name, changing lines, short I Ching reference."],
]
for row in prompt_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
set_table_width(table, [1800, 2300, 5260])
style_table(table)

add_heading(doc, "7. 最终建议", 1)
add_body(doc, "品牌命名采用 Dao Oracle，slogan 使用 Ancient Eastern wisdom for modern life decisions.")
add_body(doc, "结果页采用“天时 / 人和 / 地利”作为核心结构，英文前端命名建议为 Cosmic Timing / Human Field / Earthly Leverage。")
add_body(doc, "传统卦文和易经原始体系应作为 Ancient Root / Source Layer 放在可展开区域，增强可信度，但不压在用户第一眼理解之前。")

add_heading(doc, "8. 资料来源", 1)
sources = [
    "Co-Star: https://www.costarastrology.com/",
    "The Pattern: https://www.thepattern.com/",
    "Labyrinthos: https://labyrinthos.co/",
    "Trusted Tarot: https://www.trustedtarot.com/",
    "Grand View Research - Spiritual Wellness Apps Market: https://www.grandviewresearch.com/industry-analysis/spiritual-wellness-apps-market-report",
    "项目内部资料：01_项目总览与落地执行方案.md、02_竞品深度拆解报告.md、项目执行与运营计划书.md",
]
for source in sources:
    add_bullet(doc, source)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("Dao Oracle brand brief")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(107, 114, 128)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
